from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "START-HERE-IOS-AGENT.md"
OUTPUT = ROOT / "RADIOVAULT-IOS-AGENT-HANDOFF.docx"

INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    trpr.append(header)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_borders(table, color="D4DAE2", size="4"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def next_numbering_id(numbering, tag):
    values = []
    for node in numbering.findall(qn(tag)):
        attr = "abstractNumId" if tag == "w:abstractNum" else "numId"
        value = node.get(qn(f"w:{attr}"))
        if value is not None and value.isdigit():
            values.append(int(value))
    return max(values, default=0) + 1


def create_numbering(doc, kind):
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum")
    num_id = next_numbering_id(numbering, "w:num")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numfmt = OxmlElement("w:numFmt")
    numfmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(numfmt)
    lvltext = OxmlElement("w:lvlText")
    lvltext.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvltext)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    lvl.append(ppr)
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")
    rpr.append(rfonts)
    lvl.append(rpr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId")
    numid.set(qn("w:val"), str(num_id))
    numpr.append(ilvl)
    numpr.append(numid)


def add_inline(paragraph, text, base_size=11, base_color=None):
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_font(run, size=base_size, color=base_color)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, "Consolas", base_size - 0.5, INK)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PALE_GRAY)
            run._element.get_or_add_rPr().append(shd)
        else:
            run = paragraph.add_run(token[2:-2])
            set_font(run, size=base_size, color=base_color, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_font(run, size=base_size, color=base_color)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.08
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_GRAY)
    ppr.append(shd)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_font(run, "Consolas", 8.5, INK)
        if index < len(lines) - 1:
            run.add_break()


def add_markdown_table(doc, rows):
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[2:] if all(re.fullmatch(r":?-{3,}:?", value.strip()) for value in rows[1]) else rows[1:]
    column_count = len(header)
    if column_count == 2:
        widths = [3000, 6360]
    elif column_count == 3:
        widths = [2300, 3530, 3530]
    else:
        base = 9360 // column_count
        widths = [base] * column_count
        widths[-1] += 9360 - sum(widths)
    table = doc.add_table(rows=1, cols=column_count)
    add_table_borders(table)
    for idx, text in enumerate(header):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_inline(paragraph, text.strip(), 9.5, INK)
        for run in paragraph.runs:
            run.bold = True
    set_repeat_table_header(table.rows[0])
    for row_values in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row_values):
            paragraph = cells[idx].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_inline(paragraph, text.strip(), 9.25, None)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(doc, title, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT)
    ppr.append(shading)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "5")
        border.set(qn("w:color"), "C8D3E0")
        borders.append(border)
    ppr.append(borders)
    run = paragraph.add_run(title + "  ")
    set_font(run, size=10.5, color=DARK_BLUE, bold=True)
    add_inline(paragraph, text, 10.5, INK)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, value, end):
        run._r.append(node)
    set_font(run, size=8.5, color=MUTED)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("222222")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    left = paragraph.add_run("RADIO VAULT")
    set_font(left, size=8.5, color=BLUE, bold=True)
    right = paragraph.add_run("   |   iOS Engineering Handoff")
    set_font(right, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    prefix = fp.add_run("Radio Vault iOS Agent Handoff  |  ")
    set_font(prefix, size=8.5, color=MUTED)
    add_page_field(fp)


def add_cover(doc):
    logo = ROOT / "TheRadioVault.Desktop.Avalonia" / "Assets" / "RadioVault-Logo.png"
    top = doc.add_paragraph()
    top.paragraph_format.space_after = Pt(28)
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo.exists():
        picture = top.add_run().add_picture(str(logo), width=Inches(0.82))
        picture._inline.docPr.set("title", "Radio Vault logo")
        picture._inline.docPr.set("descr", "Radio Vault archive mark")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    run = kicker.add_run("ENGINEERING HANDOFF")
    set_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Radio Vault\niOS Client")
    set_font(run, size=30, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(34)
    run = subtitle.add_run("Goals, architecture, roadmap and implementation guardrails")
    set_font(run, size=14, color=DARK_BLUE)

    for text, bold in (
        ("Source snapshot: 0.35.0-alpha9-buildfix3", True),
        ("API v1  |  Database schema 47  |  Capability generation 40", False),
        ("Prepared 8 August 2026", False),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        set_font(run, size=10.5, color=MUTED, bold=bold)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(24)
    add_callout(
        doc,
        "Core rule",
        "The Server owns the archive and every durable operation. The iOS application is a secure native presentation client, never a second database owner.",
    )
    doc.add_page_break()


def add_contents(doc, headings):
    title = doc.add_paragraph("Contents", style="Heading 1")
    title.paragraph_format.space_before = Pt(0)
    intro = doc.add_paragraph("Use this guide together with the source files named in section 19. Executable contracts and current tests take precedence over historical notes.")
    intro.paragraph_format.space_after = Pt(10)
    num_id = create_numbering(doc, "decimal")
    for heading in headings:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        apply_numbering(p, num_id)
        clean = re.sub(r"^\d+\.\s*", "", heading)
        add_inline(p, clean, 10.5, INK)
    doc.add_page_break()


def parse_markdown(doc, text):
    lines = text.splitlines()
    headings = [line[3:].strip() for line in lines if line.startswith("## ")]
    add_contents(doc, headings)
    bullet_id = create_numbering(doc, "bullet")
    number_id = create_numbering(doc, "decimal")
    page_break_sections = {
        "3. Authoritative architecture",
        "5. Source map",
        "7. API orientation",
        "9. Transactional playback handoff",
        "12. Proposed iOS delivery roadmap",
        "14. Engineering workflow on the Mac",
        "18. Return package expectations",
    }
    index = 0
    paragraph_buffer = []
    in_code = False
    code_lines = []
    skipped_title = False
    skipped_metadata = 0

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, " ".join(part.strip() for part in paragraph_buffer), 11, None)
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if line.startswith("# ") and not skipped_title:
            skipped_title = True
            index += 1
            continue
        if skipped_title and skipped_metadata < 5 and (stripped.startswith("**Snapshot:") or stripped.startswith("**Radio Vault version:") or stripped.startswith("**Server API:") or stripped.startswith("**Database schema:") or stripped.startswith("**Connected-client")):
            skipped_metadata += 1
            index += 1
            continue
        if line.startswith("## "):
            flush_paragraph()
            heading = line[3:].strip()
            if heading in page_break_sections:
                doc.add_page_break()
            doc.add_paragraph(heading, style="Heading 1")
            index += 1
            continue
        if line.startswith("### "):
            flush_paragraph()
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
            index += 1
            continue
        if line.startswith("#### "):
            flush_paragraph()
            doc.add_paragraph(line[5:].strip(), style="Heading 3")
            index += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph()
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|") and lines[index].strip().endswith("|"):
                values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
                table_lines.append(values)
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        bullet = re.match(r"^\s*-\s+(.+)$", line)
        numbered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if bullet or numbered:
            flush_paragraph()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_numbering(paragraph, bullet_id if bullet else number_id)
            add_inline(paragraph, (bullet or numbered).group(1), 10.75, None)
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        paragraph_buffer.append(stripped)
        index += 1
    flush_paragraph()


def structural_audit(doc_path):
    import zipfile
    with zipfile.ZipFile(doc_path, "r") as archive:
        names = set(archive.namelist())
        required = {
            "[Content_Types].xml",
            "word/document.xml",
            "word/styles.xml",
            "word/numbering.xml",
            "word/settings.xml",
        }
        missing = required - names
        if missing:
            raise RuntimeError(f"DOCX structural audit failed; missing: {sorted(missing)}")
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        if "Radio Vault" not in document_xml or "Proposed iOS delivery roadmap" not in document_xml:
            raise RuntimeError("DOCX structural audit failed; expected handoff content is absent.")
        if 'w:line="300"' not in styles_xml:
            raise RuntimeError("DOCX preset audit failed; compact body line spacing is absent.")
        if 'w:left="540"' not in numbering_xml or 'w:hanging="270"' not in numbering_xml:
            raise RuntimeError("DOCX preset audit failed; list geometry is absent.")


def main():
    if not SOURCE.exists():
        raise SystemExit(f"Missing source guide: {SOURCE}")
    doc = Document()
    configure_styles(doc)
    doc.core_properties.title = "Radio Vault iOS Client Engineering Handoff"
    doc.core_properties.subject = "Goals, architecture, roadmap and implementation guardrails for a native iOS client"
    doc.core_properties.author = "Radio Vault project"
    doc.core_properties.keywords = "Radio Vault, iOS, SwiftUI, AVFoundation, engineering handoff"
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUTPUT)
    structural_audit(OUTPUT)
    print(f"Created and structurally audited: {OUTPUT}")


if __name__ == "__main__":
    main()
